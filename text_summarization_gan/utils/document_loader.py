import os
import PyPDF2
import docx

def extract_text_from_file(file_path: str) -> str:
    """Reads a direct local filepath and extracts text based on its extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def extract_text_from_pdf(file_path) -> str:
    text = ""
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    return text

def extract_text_from_docx(file_path) -> str:
    """Extracts text from paragraphs and tables in a .docx file."""
    doc = docx.Document(file_path)
    full_text = []

    # Get standard paragraph text
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    # Get data from within any tables (often used in abstracts/headers)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)

    return "\n".join(full_text)

def extract_text_from_txt(file_path) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def extract_text_from_uploaded_file(uploaded_file) -> str:
    """Helper method to parse bytes straight from the UploadedFile provided by Streamlit"""
    filename = uploaded_file.name
    ext = os.path.splitext(filename)[1].lower()
    
    if ext == '.pdf':
        reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text
    elif ext == '.docx':
        doc = docx.Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif ext == '.txt':
        return uploaded_file.getvalue().decode("utf-8")
    else:
        raise ValueError(f"Unsupported file format: {ext}")
